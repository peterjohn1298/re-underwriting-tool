"""Generate institutional investment memo as Word document."""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from config import OUTPUT_DIR

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
DARK_BLUE = RGBColor(0x2C, 0x3E, 0x6B)
GOLD = RGBColor(0xC8, 0xA9, 0x51)


def _shade(cell, color_hex):
    tc = cell._element.get_or_add_tcPr()
    tc.append(tc.makeelement(qn('w:shd'), {qn('w:fill'): color_hex, qn('w:val'): 'clear'}))


def _cur(val):
    if val is None: return "N/A"
    if abs(val) >= 1e6: return f"${val/1e6:,.1f}M"
    if abs(val) >= 1000: return f"${val:,.0f}"
    return f"${val:,.2f}"


def _pct(val):
    return f"{val*100:.2f}%" if val is not None else "N/A"


def _x(val):
    return f"{val:.2f}x" if val is not None else "N/A"


def _heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
        run.font.name = "Calibri"
    return h


def _table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        _shade(cell, "1B2A4A")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0] if p.runs else p.add_run(h)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = "Calibri"
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val) if val is not None else ""
            p = cell.paragraphs[0]
            if ci > 0: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = "Calibri"
            if ri % 2 == 0: _shade(cell, "F2F2F2")
    return t


def generate_word(pro_forma: dict, market_data: dict, job_id: str,
                  ml_valuation: dict = None,
                  lease_analysis: dict = None,
                  rent_prediction: dict = None,
                  sensitivity: dict = None,
                  backtest: dict = None,
                  monte_carlo: dict = None) -> str:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    inp = pro_forma["inputs"]
    deal = inp["deal"]
    derived = inp["derived"]
    m = pro_forma["metrics"]
    pf = pro_forma["pro_forma"]
    su = pro_forma["sources_uses"]
    rev = pro_forma["reversion"]

    # Determine which optional sections exist
    has_ml = ml_valuation and not ml_valuation.get("error")
    has_lease = lease_analysis and not lease_analysis.get("error")
    has_rent = rent_prediction and not rent_prediction.get("error")
    has_backtest = backtest and not backtest.get("error") and has_rent
    has_mc = monte_carlo and not monte_carlo.get("error")

    # ===== COVER PAGE =====
    for _ in range(6): doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("INVESTMENT MEMORANDUM")
    r.font.size = Pt(28); r.font.color.rgb = NAVY; r.bold = True; r.font.name = "Calibri"

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(derived["property_name"])
    r.font.size = Pt(22); r.font.color.rgb = GOLD; r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(deal["address"])
    r.font.size = Pt(14); r.font.color.rgb = DARK_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(deal["property_type"])
    r.font.size = Pt(13); r.font.color.rgb = DARK_BLUE

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(datetime.now().strftime("%B %d, %Y"))
    r.font.size = Pt(12); r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CONFIDENTIAL")
    r.font.size = Pt(11); r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00); r.bold = True

    doc.add_page_break()

    # ===== TABLE OF CONTENTS =====
    _heading(doc, "Table of Contents", 1)
    toc_items = [
        "1. Executive Summary",
        "2. Property Overview",
        "3. Market Analysis",
        "4. Comparable Sales",
        "5. Investment Thesis",
        "6. Financial Summary",
        "7. Risk Assessment",
        "8. Capital Structure",
    ]
    section_num = 9
    if has_ml:
        toc_items.append(f"{section_num}. ML-Based Property Valuation")
        section_num += 1
    if has_lease:
        toc_items.append(f"{section_num}. Lease Document Analysis")
        section_num += 1
    if has_rent:
        toc_items.append(f"{section_num}. Predictive Rent Growth Model")
        section_num += 1
    if has_backtest:
        toc_items.append(f"{section_num}. Model Backtest Validation")
        section_num += 1
    if has_mc:
        toc_items.append(f"{section_num}. Monte Carlo Simulation")
        section_num += 1
    toc_items.append(f"{section_num}. Next Steps")

    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
    doc.add_page_break()

    # ===== 1. EXECUTIVE SUMMARY =====
    _heading(doc, "1. Executive Summary", 1)

    irr_val = m.get("levered_irr")
    recommendation = "BUY" if irr_val and irr_val >= 0.12 else ("HOLD / CONDITIONAL" if irr_val and irr_val >= 0.08 else "PASS")

    # Deal snapshot
    _table(doc, ["Attribute", "Details"], [
        ["Property", derived["property_name"]],
        ["Type", deal["property_type"]],
        ["Address", deal["address"]],
        ["Year Built", str(deal["year_built"])],
        ["Units / SF", f"{deal['total_units']} units / {deal['total_sf']:,.0f} SF"],
        ["Purchase Price", _cur(deal["purchase_price"])],
        ["Price / Unit", _cur(derived["price_per_unit"])],
        ["Current NOI", _cur(deal["current_noi"])],
        ["In-Place Rent", f"${deal['in_place_rent']:,.0f}/mo avg"],
        ["Market Rent", f"${deal['market_rent']:,.0f}/mo avg"],
        ["Occupancy", f"{deal['occupancy']*100:.0f}%"],
        ["Total CapEx", _cur(derived["total_capex"])],
        ["Hold Period", f"{deal['hold_period_years']} years"],
    ])
    doc.add_paragraph()

    # Returns
    _heading(doc, "Return Metrics", 2)
    ret_rows = [
        ["Levered IRR", _pct(m.get("levered_irr"))],
        ["After-Tax IRR", _pct(m.get("after_tax_irr"))],
        ["Unlevered IRR", _pct(m.get("unlevered_irr"))],
        ["Equity Multiple", _x(m.get("equity_multiple"))],
        ["After-Tax Equity Multiple", _x(m.get("after_tax_equity_multiple"))],
        ["Cash-on-Cash (Year 1)", _pct(m.get("cash_on_cash_yr1"))],
        ["After-Tax CoC (Year 1)", _pct(m.get("cash_on_cash_yr1_after_tax"))],
        ["DSCR (Year 1)", _x(m.get("dscr_yr1"))],
        ["Stabilized DSCR", _x(m.get("stabilized_dscr"))],
        ["Going-In Cap Rate", _pct(m.get("going_in_cap_rate"))],
        ["Exit Cap Rate", _pct(m.get("exit_cap_rate"))],
        ["Yield on Cost", _pct(m.get("yield_on_cost"))],
        ["Stabilized YOC", _pct(m.get("stabilized_yoc"))],
    ]
    if m.get("used_variable_growth"):
        ret_rows.append(["Growth Model", "Variable (ML-predicted year-by-year rates)"])
    _table(doc, ["Metric", "Value"], ret_rows)
    doc.add_paragraph()

    rec_color = RGBColor(0x00, 0x80, 0x00) if recommendation == "BUY" else (
        RGBColor(0xFF, 0x99, 0x00) if "HOLD" in recommendation else RGBColor(0xCC, 0x00, 0x00))
    p = doc.add_paragraph()
    p.add_run("Investment Recommendation: ").bold = True
    r = p.add_run(recommendation)
    r.font.color.rgb = rec_color; r.bold = True; r.font.size = Pt(16)

    if irr_val:
        rationale = (
            f"The investment projects a levered IRR of {_pct(irr_val)} and an equity multiple of "
            f"{_x(m.get('equity_multiple'))} over a {deal['hold_period_years']}-year hold. "
        )
        if derived["rent_premium_potential"] > 0:
            rationale += (
                f"Significant value-add upside exists with ${derived['rent_premium_potential']:,.0f}/unit/mo "
                f"rent premium potential (in-place ${deal['in_place_rent']:,.0f} vs. market ${deal['market_rent']:,.0f})."
            )
        doc.add_paragraph(rationale)
    doc.add_page_break()

    # ===== 2. PROPERTY OVERVIEW =====
    _heading(doc, "2. Property Overview", 1)
    doc.add_paragraph(
        f"{derived['property_name']} is a {deal['total_units']}-unit {deal['property_type'].lower()} property "
        f"located at {deal['address']}. Built in {deal['year_built']}, the property comprises approximately "
        f"{deal['total_sf']:,.0f} square feet ({deal['total_sf']/deal['total_units']:,.0f} SF/unit avg). "
        f"Current occupancy is {deal['occupancy']*100:.0f}% with in-place rents averaging ${deal['in_place_rent']:,.0f}/month "
        f"versus market rents of ${deal['market_rent']:,.0f}/month."
    )

    if deal.get("deferred_maintenance", 0) > 0 or deal.get("planned_capex", 0) > 0:
        doc.add_paragraph()
        _heading(doc, "Capital Expenditure Plan", 2)
        if deal["deferred_maintenance"] > 0:
            doc.add_paragraph(f"Deferred Maintenance: {_cur(deal['deferred_maintenance'])}", style="List Bullet")
        if deal["planned_capex"] > 0:
            desc = deal.get("capex_description", "unit upgrades")
            doc.add_paragraph(f"Planned CapEx ({desc}): {_cur(deal['planned_capex'])} "
                              f"({_cur(derived['capex_per_unit'])}/unit)", style="List Bullet")
        doc.add_paragraph(f"Total CapEx Budget: {_cur(derived['total_capex'])}", style="List Bullet")
    doc.add_page_break()

    # ===== 3. MARKET ANALYSIS =====
    _heading(doc, "3. Market Analysis", 1)
    demo = market_data.get("demographics", {})
    doc.add_paragraph(demo.get("summary", f"{derived['city']}, {derived['state']} is a growing market."))

    _heading(doc, "Market Indicators", 2)
    cap_data = market_data.get("cap_rates", {})
    rent_data = market_data.get("rent_trends", {})
    structured = demo.get("structured", {})

    indicator_rows = [
        ["Market Cap Rate", f"{cap_data.get('average_cap_rate', 'N/A')}%" if isinstance(cap_data.get('average_cap_rate'), (int, float)) else "N/A"],
        ["Avg Rent Growth (CPI Shelter)", f"{rent_data.get('average_growth', 'N/A')}%" if isinstance(rent_data.get('average_growth'), (int, float)) else "N/A"],
    ]

    # Add structured data from APIs
    if structured.get("population"):
        indicator_rows.append(["State Population", f"{structured['population']:,}"])
    if structured.get("median_income"):
        indicator_rows.append(["Median Household Income", f"${structured['median_income']:,}"])
    if structured.get("median_rent"):
        indicator_rows.append(["Median Gross Rent", f"${structured['median_rent']:,}/mo"])
    if structured.get("unemployment_rate"):
        indicator_rows.append(["Unemployment Rate", f"{structured['unemployment_rate']}%"])
    if structured.get("vacancy_rate"):
        indicator_rows.append(["Housing Vacancy Rate", f"{structured['vacancy_rate']}%"])

    indicator_rows.append(["Cap Rate Source", cap_data.get("source", "defaults")])
    indicator_rows.append(["Demographics Source", demo.get("source", "defaults")])
    indicator_rows.append(["Rent Trends Source", rent_data.get("source", "defaults")])

    _table(doc, ["Indicator", "Value"], indicator_rows)

    if cap_data.get("treasury_10yr"):
        doc.add_paragraph()
        doc.add_paragraph(
            f"Cap rate derived from 10-Year Treasury ({cap_data['treasury_10yr']}%) plus "
            f"{cap_data.get('spread_used', 2.0)}% property-type spread. "
            f"Data sourced from Federal Reserve Economic Data (FRED), Census Bureau ACS, and Bureau of Labor Statistics.",
        )

    search_results = demo.get("search_results", [])
    if search_results:
        doc.add_paragraph()
        _heading(doc, "Research Sources", 3)
        for sr in search_results[:5]:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(sr.get("source", "")); r.bold = True
            if sr.get("snippet"):
                p.add_run(f": {sr['snippet'][:200]}")
    doc.add_page_break()

    # ===== 4. COMPARABLE SALES =====
    _heading(doc, "4. Comparable Sales", 1)
    comps = market_data.get("comps", {}).get("comps", [])
    if comps and isinstance(comps[0], dict) and "price_per_unit" in comps[0]:
        rows = []
        for i, c in enumerate(comps[:7], 1):
            rows.append([
                str(i), c.get("name", f"Comp {i}"), _cur(c.get("price_per_unit")),
                f"{c.get('cap_rate', 'N/A')}%", str(c.get("year", "")), str(c.get("units", "")),
            ])
        _table(doc, ["#", "Property", "Price/Unit", "Cap Rate", "Year", "Units"], rows)
        ppus = [c["price_per_unit"] for c in comps if isinstance(c.get("price_per_unit"), (int, float))]
        if ppus:
            doc.add_paragraph()
            subj_ppu = deal["purchase_price"] / deal["total_units"] if deal["total_units"] > 0 else 0
            avg_ppu = sum(ppus) / len(ppus)
            doc.add_paragraph(
                f"Average comp price per unit is {_cur(avg_ppu)} versus the subject at {_cur(subj_ppu)}/unit. "
                f"The subject {'trades at a discount to' if subj_ppu < avg_ppu else 'is in line with'} comparable transactions."
            )
    else:
        doc.add_paragraph("See Excel workbook for detailed comparable sales analysis.")
    doc.add_page_break()

    # ===== 5. INVESTMENT THESIS =====
    _heading(doc, "5. Investment Thesis", 1)

    _heading(doc, "Key Return Drivers", 2)
    drivers = [
        f"Going-in cap rate of {_pct(m.get('going_in_cap_rate'))} on current NOI of {_cur(deal['current_noi'])}.",
    ]
    if derived["rent_premium_potential"] > 0:
        drivers.append(
            f"Value-add opportunity: ${derived['rent_premium_potential']:,.0f}/unit/mo rent upside "
            f"(${deal['in_place_rent']:,.0f} in-place -> ${deal['market_rent']:,.0f} market) "
            f"representing {_cur(derived['rent_premium_potential'] * deal['total_units'] * 12)} in annual revenue upside."
        )
    drivers.append(
        f"Conservative {_pct(deal['revenue_growth_rate'])} annual revenue growth assumption."
    )
    drivers.append(
        f"Leverage at {_pct(deal['ltv'])} LTV / {_pct(deal['interest_rate'])} provides positive leverage "
        f"with {_x(m.get('dscr_yr1'))} DSCR."
    )
    for d in drivers:
        doc.add_paragraph(d, style="List Bullet")

    if derived["total_capex"] > 0:
        doc.add_paragraph()
        _heading(doc, "Value-Add Strategy", 2)
        desc = deal.get("capex_description", "property improvements")
        doc.add_paragraph(
            f"Total capital budget of {_cur(derived['total_capex'])} ({_cur(derived['capex_per_unit'])}/unit) "
            f"for {desc}. The investment targets rent premiums of ${derived['rent_premium_potential']:,.0f}/unit/month "
            f"post-renovation, driving stabilized yield on cost to {_pct(m.get('stabilized_yoc'))}."
        )
    doc.add_page_break()

    # ===== 6. FINANCIAL SUMMARY =====
    _heading(doc, "6. Financial Summary", 1)

    _heading(doc, "Key Assumptions", 2)
    _table(doc, ["Assumption", "Value"], [
        ["Purchase Price", _cur(deal["purchase_price"])],
        ["Total Project Cost", _cur(derived["total_project_cost"])],
        ["LTV / Rate", f"{_pct(deal['ltv'])} / {_pct(deal['interest_rate'])}"],
        ["Amortization", f"{deal['amortization_years']} years"],
        ["Revenue Growth", _pct(deal["revenue_growth_rate"])],
        ["Expense Growth", _pct(deal["expense_growth_rate"])],
        ["Exit Cap Rate", _pct(m.get("exit_cap_rate"))],
        ["Hold Period", f"{deal['hold_period_years']} years"],
    ])
    doc.add_paragraph()

    # NOI trajectory
    _heading(doc, "NOI & Cash Flow Trajectory", 2)
    noi_rows = []
    for yr in pf[:deal["hold_period_years"]]:
        row = [
            f"Year {yr['year']}", f"${yr['rent_per_unit']:,.0f}",
            f"{yr['occupancy']*100:.0f}%", _cur(yr["noi"]), _cur(yr["btcf"]),
        ]
        if "atcf" in yr:
            row.append(_cur(yr["atcf"]))
        noi_rows.append(row)
    noi_headers = ["Year", "Rent/Unit", "Occ.", "NOI", "BTCF"]
    if pf and "atcf" in pf[0]:
        noi_headers.append("ATCF")
    _table(doc, noi_headers, noi_rows)
    doc.add_paragraph()

    # Exit
    _heading(doc, "Exit Summary", 2)
    exit_rows = [
        ["Exit Year", str(rev["exit_year"])],
        ["Forward NOI", _cur(rev["forward_noi"])],
        ["Exit Cap Rate", _pct(rev["exit_cap_rate"])],
        ["Gross Sale Price", _cur(rev["sale_price"])],
        ["Sale Costs", _cur(rev["sale_costs"])],
        ["Loan Payoff", _cur(rev["loan_balance"])],
        ["Net Sale Proceeds (Pre-Tax)", _cur(rev["net_sale_proceeds"])],
    ]
    if rev.get("total_tax_on_sale"):
        exit_rows.extend([
            ["Depreciation Recapture Tax", f"({_cur(rev['depreciation_recapture_tax'])})"],
            ["Capital Gains Tax", f"({_cur(rev['capital_gains_tax'])})"],
            ["Net Sale Proceeds (After-Tax)", _cur(rev["net_sale_proceeds_after_tax"])],
        ])
    _table(doc, ["Item", "Value"], exit_rows)
    doc.add_page_break()

    # ===== 7. RISK ASSESSMENT =====
    _heading(doc, "7. Risk Assessment", 1)
    dscr = m.get("dscr_yr1", 0)
    _table(doc, ["Risk Category", "Severity", "Description"], [
        ["Market Risk", "Medium", "Economic downturn or oversupply affecting occupancy and rents"],
        ["Interest Rate Risk", "High" if deal["interest_rate"] > 0.07 else "Medium",
         f"Current rate {_pct(deal['interest_rate'])}; refinancing risk at maturity"],
        ["Execution Risk",
         "Medium" if derived["total_capex"] > 0 else "Low",
         f"CapEx execution on {_cur(derived['total_capex'])} renovation budget" if derived["total_capex"] > 0
         else "Minimal renovation risk"],
        ["Operational Risk", "Low" if dscr > 1.3 else "Medium",
         f"DSCR of {_x(dscr)} provides {'adequate' if dscr > 1.25 else 'thin'} coverage"],
        ["Liquidity Risk", "Medium", "Real estate illiquidity; exit timing uncertain"],
    ])
    doc.add_paragraph()
    doc.add_paragraph(
        "Key mitigants include conservative underwriting, adequate debt service coverage, "
        "and a diversified unit base. The value-add strategy provides multiple return drivers "
        "beyond market appreciation alone."
    )
    doc.add_page_break()

    # ===== 8. CAPITAL STRUCTURE =====
    _heading(doc, "8. Capital Structure", 1)

    _heading(doc, "Sources & Uses", 2)
    total_s = su["sources"]["total"]
    _table(doc, ["Sources", "Amount", "% of Total"], [
        ["Senior Debt", _cur(su["sources"]["senior_debt"]),
         _pct(su["sources"]["senior_debt"]/total_s) if total_s else "0%"],
        ["Sponsor Equity", _cur(su["sources"]["sponsor_equity"]),
         _pct(su["sources"]["sponsor_equity"]/total_s) if total_s else "0%"],
        ["Total", _cur(total_s), "100.00%"],
    ])
    doc.add_paragraph()

    total_u = su["uses"]["total"]
    uses_rows = [
        ["Purchase Price", _cur(su["uses"]["purchase_price"]),
         _pct(su["uses"]["purchase_price"]/total_u) if total_u else "0%"],
        ["Closing Costs", _cur(su["uses"]["closing_costs"]),
         _pct(su["uses"]["closing_costs"]/total_u) if total_u else "0%"],
    ]
    if su["uses"]["deferred_maintenance"] > 0:
        uses_rows.append(["Deferred Maintenance", _cur(su["uses"]["deferred_maintenance"]),
                          _pct(su["uses"]["deferred_maintenance"]/total_u)])
    if su["uses"]["planned_capex"] > 0:
        uses_rows.append([f"Planned CapEx ({deal.get('capex_description', '')})",
                          _cur(su["uses"]["planned_capex"]),
                          _pct(su["uses"]["planned_capex"]/total_u)])
    uses_rows.append(["Total", _cur(total_u), "100.00%"])
    _table(doc, ["Uses", "Amount", "% of Total"], uses_rows)
    doc.add_paragraph()

    _heading(doc, "Debt Terms", 2)
    _table(doc, ["Term", "Details"], [
        ["Loan Amount", _cur(su["sources"]["senior_debt"])],
        ["LTV", _pct(deal["ltv"])],
        ["Interest Rate", _pct(deal["interest_rate"])],
        ["Amortization", f"{deal['amortization_years']} years"],
        ["Loan Term", f"{deal['loan_term_years']} years"],
        ["IO Period", f"{deal['io_period_years']} years"],
    ])
    doc.add_page_break()

    # ===== OPTIONAL: ML VALUATION =====
    section_num = 9
    if has_ml:
        _heading(doc, f"{section_num}. ML-Based Property Valuation", 1)
        section_num += 1

        doc.add_paragraph(
            f"A GradientBoosting regression model was trained on {ml_valuation['training_samples']} "
            f"SYNTHETIC property records calibrated to current FRED/Census macro indicators. "
            f"An 80/20 train/test split yielded a held-out test R² of {ml_valuation.get('test_r2', ml_valuation['r2_score']):.4f} "
            f"with a test MAPE of {ml_valuation.get('test_mape', 'N/A')}%. "
            f"The model uses {ml_valuation['features_used']} features."
        )
        doc.add_paragraph()

        assessment = ml_valuation["assessment"]
        p = doc.add_paragraph()
        p.add_run("ML Assessment: ").bold = True
        color = RGBColor(0x00, 0x80, 0x00) if assessment == "UNDERVALUED" else (
            RGBColor(0xCC, 0x00, 0x00) if assessment == "OVERVALUED" else NAVY)
        r = p.add_run(assessment)
        r.font.color.rgb = color; r.bold = True; r.font.size = Pt(14)
        doc.add_paragraph()

        _table(doc, ["Metric", "Value"], [
            ["Predicted Value/Unit", _cur(ml_valuation["predicted_value_per_unit"])],
            ["Predicted Total Value", _cur(ml_valuation["predicted_total_value"])],
            ["Actual Price/Unit", _cur(ml_valuation["actual_price_per_unit"])],
            ["Premium/Discount", f"{ml_valuation['premium_discount_pct']:.1f}%"],
            ["Train R²", f"{ml_valuation.get('train_r2', 'N/A')}"],
            ["Test R² (held-out)", f"{ml_valuation.get('test_r2', ml_valuation['r2_score']):.4f}"],
            ["Test MAE", _cur(ml_valuation.get("test_mae"))],
            ["Test MAPE", f"{ml_valuation.get('test_mape', 'N/A')}%"],
            ["Model Type", ml_valuation["model_type"]],
            ["Training Data", "Synthetic (calibrated to FRED/Census)"],
        ])
        doc.add_paragraph()

        # Top features
        _heading(doc, "Key Value Drivers (Feature Importance)", 2)
        top_features = list(ml_valuation["feature_importances"].items())[:8]
        feat_rows = [[feat.replace("_", " ").title(), f"{imp:.4f}"] for feat, imp in top_features]
        _table(doc, ["Feature", "Importance"], feat_rows)
        doc.add_paragraph()

        doc.add_paragraph(
            "Note: The ML model uses synthetic training data calibrated to real macroeconomic indicators. "
            "Results should be used as a supplemental data point alongside traditional appraisal methods.",
        ).runs[0].font.italic = True
        doc.add_page_break()

    # ===== OPTIONAL: LEASE ANALYSIS =====
    if has_lease:
        _heading(doc, f"{section_num}. Lease Document Analysis", 1)
        section_num += 1

        if lease_analysis.get("summary"):
            doc.add_paragraph(lease_analysis["summary"])
            doc.add_paragraph()

        _heading(doc, "Extracted Lease Terms", 2)
        term_rows = []
        term_fields = [
            ("Tenant", "tenant_name"), ("Landlord", "landlord_name"),
            ("Lease Type", "lease_type"), ("Monthly Rent", "monthly_rent"),
            ("Annual Rent", "annual_rent"), ("Term (months)", "lease_term_months"),
            ("Start Date", "lease_start_date"), ("End Date", "lease_end_date"),
            ("Escalation", "escalation_clause"), ("Annual Escalation", "annual_escalation_pct"),
            ("Renewal Options", "renewal_options"), ("Security Deposit", "security_deposit"),
            ("TI Allowance", "ti_allowance"), ("CAM Charges", "cam_charges"),
        ]
        for label, key in term_fields:
            val = lease_analysis.get(key)
            if val is not None:
                if isinstance(val, (int, float)) and val > 100:
                    term_rows.append([label, _cur(val)])
                elif isinstance(val, float):
                    term_rows.append([label, f"{val:.2f}%"])
                else:
                    term_rows.append([label, str(val)])

        if term_rows:
            _table(doc, ["Term", "Value"], term_rows)
            doc.add_paragraph()

        # Key clauses
        clauses = lease_analysis.get("key_clauses", [])
        if clauses:
            _heading(doc, "Key Clauses Identified", 2)
            for clause in clauses:
                doc.add_paragraph(clause, style="List Bullet")
            doc.add_paragraph()

        # Risk flags
        flags = lease_analysis.get("risk_flags", [])
        if flags:
            _heading(doc, "Risk Flags", 2)
            for flag in flags:
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(flag)
                r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            doc.add_paragraph()

        method = lease_analysis.get("analysis_method", "N/A")
        p = doc.add_paragraph(f"Analysis method: {method}")
        p.runs[0].font.italic = True
        p.runs[0].font.size = Pt(9)
        doc.add_page_break()

    # ===== OPTIONAL: RENT PREDICTION =====
    if has_rent:
        _heading(doc, f"{section_num}. Predictive Rent Growth Model", 1)
        section_num += 1

        doc.add_paragraph(
            f"A polynomial regression model (degree=2) was trained on {rent_prediction['training_points']} "
            f"historical CPI Shelter growth observations from FRED. The model forecasts annual rent growth "
            f"rates for the {rent_prediction['hold_period']}-year hold period, with predictions clamped to [-2%, +8%]."
        )
        doc.add_paragraph()

        _table(doc, ["Metric", "Value"], [
            ["Method", rent_prediction["method"]],
            ["Data Source", rent_prediction["data_source"]],
            ["Historical Avg Growth", f"{rent_prediction['historical_avg']}%"],
            ["Predicted Avg Growth", f"{rent_prediction['avg_predicted_growth']}%"],
            ["Current Rent/Unit", _cur(rent_prediction["current_rent"])],
        ])
        doc.add_paragraph()

        _heading(doc, "Rent Growth Forecast", 2)
        forecast_rows = []
        for i in range(len(rent_prediction["predicted_rates"])):
            forecast_rows.append([
                f"Year {i+1}",
                f"{rent_prediction['predicted_rates'][i]:.2f}%",
                f"${rent_prediction['predicted_rents_per_unit'][i]:,.0f}",
                _cur(rent_prediction["predicted_annual_revenue"][i]),
            ])
        _table(doc, ["Year", "Growth Rate", "Rent/Unit", "Annual Revenue"], forecast_rows)
        doc.add_paragraph()

        if rent_prediction.get("historical_rates"):
            _heading(doc, "Historical Growth Rates", 2)
            doc.add_paragraph(
                f"Recent historical CPI Shelter growth rates: "
                f"{', '.join(f'{r:.1f}%' for r in rent_prediction['historical_rates'][-5:])}"
            )
        doc.add_page_break()

    # ===== OPTIONAL: BACKTEST VALIDATION =====
    if has_backtest:
        _heading(doc, f"{section_num}. Model Backtest Validation", 1)
        section_num += 1

        doc.add_paragraph(
            f"The rent prediction model was backtested using a {backtest['split_pct']}% / "
            f"{100 - backtest['split_pct']}% time-based train/test split on "
            f"{backtest['total_periods']} historical CPI Shelter growth observations."
        )
        doc.add_paragraph()

        bt_rows = [
            ["Train Periods", str(backtest["train_periods"])],
            ["Test Periods", str(backtest["test_periods"])],
            ["MAE", f"{backtest['mae']:.3f}%"],
            ["RMSE", f"{backtest['rmse']:.3f}%"],
            ["Direction Accuracy", f"{backtest['direction_accuracy']}%" if backtest.get("direction_accuracy") else "N/A"],
            ["Mean Actual Growth", f"{backtest['mean_actual']:.2f}%"],
            ["Mean Predicted Growth", f"{backtest['mean_predicted']:.2f}%"],
            ["Quality Assessment", backtest.get("quality", "N/A")],
        ]
        if backtest.get("has_blended"):
            bt_rows.append(["Blended MAE (ZORI+CPI)", f"{backtest['blended_mae']:.3f}%"])
            bt_rows.append(["Blended RMSE (ZORI+CPI)", f"{backtest['blended_rmse']:.3f}%"])

        _table(doc, ["Metric", "Value"], bt_rows)
        doc.add_paragraph()

        # Actual vs Predicted comparison
        if backtest.get("actual_values") and backtest.get("predicted_values"):
            _heading(doc, "Actual vs Predicted", 2)
            comp_rows = []
            for i in range(len(backtest["actual_values"])):
                date_label = backtest["test_dates"][i] if i < len(backtest.get("test_dates", [])) else f"T+{i+1}"
                comp_rows.append([
                    date_label,
                    f"{backtest['actual_values'][i]:.2f}%",
                    f"{backtest['predicted_values'][i]:.2f}%",
                    f"{backtest['actual_values'][i] - backtest['predicted_values'][i]:.3f}%",
                ])
            _table(doc, ["Period", "Actual", "Predicted", "Error"], comp_rows)
        doc.add_page_break()

    # ===== OPTIONAL: MONTE CARLO SIMULATION =====
    if has_mc:
        _heading(doc, f"{section_num}. Monte Carlo Simulation", 1)
        section_num += 1

        doc.add_paragraph(
            f"A Monte Carlo simulation with {monte_carlo.get('n_iterations', 1000):,} iterations was run "
            f"to assess the probability distribution of investment returns. Key input parameters were "
            f"randomized within realistic ranges."
        )
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run("Key Finding: ").bold = True
        p.add_run(monte_carlo.get("summary", ""))
        doc.add_paragraph()

        # Percentiles
        _heading(doc, "IRR Percentiles", 2)
        pctile_rows = [[k, f"{v:.1f}%"] for k, v in monte_carlo.get("percentiles", {}).items()]
        if pctile_rows:
            _table(doc, ["Percentile", "IRR"], pctile_rows)
        doc.add_paragraph()

        # Probabilities
        _heading(doc, "Probability Analysis", 2)
        prob_rows = [[k, f"{v:.1f}%"] for k, v in monte_carlo.get("probabilities", {}).items()]
        if prob_rows:
            _table(doc, ["Threshold", "Probability"], prob_rows)
        doc.add_page_break()

    # ===== NEXT STEPS =====
    _heading(doc, f"{section_num}. Next Steps", 1)
    _heading(doc, "Due Diligence Checklist", 2)
    for item in [
        "Property inspection and condition assessment",
        "Phase I Environmental Site Assessment",
        "Title search and review",
        "Rent roll verification and tenant interviews",
        "Historical financial statement review (3 years)",
        "Tax assessment review and appeal analysis",
        "Insurance coverage evaluation",
        "Market study and independent appraisal",
        "Loan commitment and legal review",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    _heading(doc, "Disclaimer", 2)
    p = doc.add_paragraph(
        "This investment memorandum is provided for informational purposes only and does not "
        "constitute an offer to sell or a solicitation of an offer to buy any security. Financial "
        "projections are based on assumptions that may not be realized. Actual results may differ "
        "materially. Past performance is not indicative of future results. Investors should conduct "
        "their own due diligence and consult with their advisors before making any investment decision."
    )
    p.runs[0].font.size = Pt(9)
    p.runs[0].font.italic = True

    if has_ml or has_lease or has_rent:
        doc.add_paragraph()
        p = doc.add_paragraph(
            "AI/ML Disclaimer: Machine learning valuations, NLP lease analysis, and predictive rent "
            "models use synthetic training data and statistical methods that carry inherent limitations. "
            "These outputs should supplement, not replace, professional judgment and independent appraisals."
        )
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.italic = True

    # Header / footer
    section = doc.sections[0]
    header_p = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
    header_p.text = f"{derived['property_name']} | Investment Memorandum"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header_p.runs:
        run.font.size = Pt(8); run.font.color.rgb = NAVY

    footer_p = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    footer_p.text = "Confidential"
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_p.runs:
        run.font.size = Pt(8); run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    filename = f"investment_memo_{job_id}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)
    doc.save(filepath)
    return filepath
