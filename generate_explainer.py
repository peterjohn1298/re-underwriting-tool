"""Generate an explanatory Word document about the RE Underwriting Tool."""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Colors
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
GOLD = RGBColor(0xC8, 0xA9, 0x51)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY = RGBColor(0x66, 0x66, 0x66)
GREEN = RGBColor(0x28, 0xA7, 0x45)

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "outputs", "RE_Underwriting_Tool_Explainer.docx")
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)


def _shade_cell(cell, color_hex):
    tc = cell._element.get_or_add_tcPr()
    tc.append(tc.makeelement(qn('w:shd'), {qn('w:fill'): color_hex, qn('w:val'): 'clear'}))


def _add_heading(doc, text, level=1, color=NAVY):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h


def _add_para(doc, text, bold=False, italic=False, size=11, color=DARK_GRAY, spacing_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(spacing_after)
    return p


def _add_bullet(doc, text, bold_prefix="", size=11):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(size)
        run.font.color.rgb = DARK_GRAY
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = DARK_GRAY
    else:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = DARK_GRAY
    return p


def _add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


def generate():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ==================== COVER PAGE ====================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RE Underwriting Tool")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = NAVY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("AI-Powered Real Estate Investment Analysis Platform")
    run.font.size = Pt(18)
    run.font.color.rgb = GOLD

    doc.add_paragraph()

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run("How This Tool Works, Why It Matters,\nand How AI/ML Makes It Institutional-Grade")
    run.font.size = Pt(14)
    run.font.color.rgb = LIGHT_GRAY

    doc.add_paragraph()
    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run("February 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = LIGHT_GRAY

    doc.add_page_break()

    # ==================== TABLE OF CONTENTS ====================
    _add_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. What Is This Tool?",
        "2. The Problem It Solves",
        "3. How It Works — Step by Step",
        "4. AI & Machine Learning Integration — The Core Innovation",
        "   4.1  ML Property Valuation (GradientBoosting Regression)",
        "   4.2  Predictive Rent Growth (Polynomial Regression + Zillow ZORI Blending)",
        "   4.3  Historical Backtest Validation",
        "   4.4  Monte Carlo Simulation (1,000 Scenario Analysis)",
        "   4.5  Natural Language Processing for Lease Analysis",
        "   4.6  Multi-Signal AI Recommendation Engine",
        "   4.7  Interactive What-If Analysis",
        "5. Real-World Data Integration",
        "6. What the Tool Produces",
        "7. Who Benefits From This Tool?",
        "8. Technical Architecture Summary",
        "9. Why This Matters — The Bigger Picture",
    ]
    for item in toc_items:
        _add_para(doc, item, size=11, spacing_after=3)

    doc.add_page_break()

    # ==================== SECTION 1 ====================
    _add_heading(doc, "1. What Is This Tool?", level=1)

    _add_para(doc,
        "The RE Underwriting Tool is a web-based application designed to analyze real estate investment "
        "opportunities with the rigor and sophistication of a professional investment firm — but accessible "
        "to anyone with a web browser. Think of it as having a team of financial analysts, data scientists, "
        "and market researchers all working together instantly to evaluate whether a property is worth buying."
    )

    _add_para(doc,
        "In the real estate industry, \"underwriting\" means carefully analyzing a property's financials to "
        "determine whether it's a good investment. This process traditionally requires expensive software "
        "(like ARGUS Enterprise, which costs thousands of dollars per year), teams of analysts, and weeks "
        "of work. Our tool does this in under 30 seconds, completely free, and with AI-powered insights "
        "that go beyond what traditional tools offer."
    )

    _add_para(doc,
        "The tool combines traditional financial modeling with cutting-edge artificial intelligence and "
        "machine learning to provide a comprehensive, data-driven investment recommendation. It doesn't "
        "just crunch numbers — it learns from data, predicts future trends, simulates thousands of "
        "possible outcomes, and even reads legal documents to extract critical information."
    )

    doc.add_page_break()

    # ==================== SECTION 2 ====================
    _add_heading(doc, "2. The Problem It Solves", level=1)

    _add_para(doc,
        "Imagine you're considering buying an apartment building for $12 million. Before committing that "
        "kind of money, you need answers to critical questions:"
    )

    bullets = [
        "Will I make money? What's my expected return over 7 years?",
        "What if rents don't grow as fast as I expect?",
        "What if interest rates change, or the economy slows down?",
        "Is $12 million a fair price, or am I overpaying?",
        "What does the lease actually say? Are there hidden risks?",
        "How confident should I be in these projections?",
        "How does this deal compare to other opportunities I'm evaluating?",
    ]
    for b in bullets:
        _add_bullet(doc, b)

    _add_para(doc, "")
    _add_para(doc,
        "Traditionally, answering these questions requires hiring consultants, building complex Excel "
        "spreadsheets from scratch, manually researching market data, and hoping you haven't made "
        "a formula error somewhere. This tool automates ALL of it and adds AI-driven insights that "
        "would be impossible to produce manually."
    )

    doc.add_page_break()

    # ==================== SECTION 3 ====================
    _add_heading(doc, "3. How It Works — Step by Step", level=1)

    _add_para(doc, "Here's what happens when you use the tool, explained in plain language:", bold=True)
    _add_para(doc, "")

    _add_heading(doc, "Step 1: You Enter the Deal Details", level=2)
    _add_para(doc,
        "You fill in a simple form with the property information: the address, purchase price, "
        "number of apartments, current rent levels, occupancy rate, and financing terms. This is "
        "the same information any buyer would gather when evaluating a property."
    )

    _add_heading(doc, "Step 2: The Tool Fetches Live Economic Data", level=2)
    _add_para(doc,
        "The tool automatically connects to three U.S. government data sources (FRED, Census Bureau, "
        "and Bureau of Labor Statistics) to pull real-time economic indicators for the property's "
        "specific market — things like local median income, population growth, unemployment rates, "
        "mortgage rates, inflation data, and housing market statistics. This ensures the analysis is "
        "grounded in actual market conditions, not assumptions."
    )

    _add_heading(doc, "Step 3: AI/ML Models Analyze the Data", level=2)
    _add_para(doc,
        "This is where the magic happens. Six different AI and machine learning systems activate "
        "simultaneously, each handling a different aspect of the analysis. (More detail in Section 4.)"
    )

    _add_heading(doc, "Step 4: Financial Model Runs", level=2)
    _add_para(doc,
        "A complete 10-year financial projection (called a \"pro forma\") is built, calculating "
        "year-by-year rental income, operating expenses, debt service payments, cash flows, taxes, "
        "depreciation, and the eventual sale of the property. If the AI rent predictor is enabled, "
        "the model uses AI-predicted rent growth rates instead of flat assumptions — making the "
        "projections far more realistic."
    )

    _add_heading(doc, "Step 5: Results Dashboard", level=2)
    _add_para(doc,
        "Everything is presented on a rich, interactive dashboard with charts, tables, color-coded "
        "recommendations, and downloadable reports (Excel workbook, Word investment memo, and PDF report). "
        "The dashboard includes interactive sliders that let you change assumptions and see the impact "
        "on returns in real-time."
    )

    doc.add_page_break()

    # ==================== SECTION 4 ====================
    _add_heading(doc, "4. AI & Machine Learning Integration — The Core Innovation", level=1)

    _add_para(doc,
        "This is the heart of what makes this tool special. While traditional underwriting tools "
        "rely purely on static formulas and manual assumptions, this tool incorporates seven distinct "
        "AI/ML systems that work together to produce institutional-quality analysis. Let's break each one down.",
        bold=True
    )

    doc.add_page_break()

    # --- 4.1 ---
    _add_heading(doc, "4.1  ML Property Valuation (GradientBoosting Regression)", level=2)

    _add_para(doc, "What it does:", bold=True)
    _add_para(doc,
        "This model answers the fundamental question: \"Is the asking price fair?\" It uses a machine "
        "learning algorithm called GradientBoosting Regression to predict what a property SHOULD be "
        "worth based on its characteristics and the current market environment."
    )

    _add_para(doc, "How it works in simple terms:", bold=True)
    _add_para(doc,
        "Imagine you wanted to guess the price of a house. You'd look at factors like size, location, "
        "age, and condition. A human might juggle 4-5 factors in their head. Our ML model simultaneously "
        "analyzes 18 different factors — including economic indicators most humans wouldn't think to check."
    )

    _add_para(doc, "The 18 features the model uses:", bold=True)

    _add_table(doc,
        ["Category", "Features", "Why It Matters"],
        [
            ["Property", "Units, SF/unit, year built, occupancy,\nin-place rent, market rent, NOI/unit,\nproperty class", "Core physical and financial characteristics"],
            ["Market", "Cap rate, median income, population,\nunemployment rate", "Local economic health and demand"],
            ["Macro Economy", "Mortgage rate, CPI inflation, housing\nstarts, rental vacancy, treasury spread,\nmedian Census rent", "National economic forces that affect\nall real estate values"],
        ]
    )

    _add_para(doc, "")
    _add_para(doc, "Technical details:", bold=True)

    _add_bullet(doc, " The model trains on 800 synthetic data points calibrated to real FRED/Census/BLS macro indicators", bold_prefix="Training Data:")
    _add_bullet(doc, " GradientBoosting with 250 estimators, max_depth=4, learning_rate=0.1", bold_prefix="Algorithm:")
    _add_bullet(doc, " Honest 80/20 train/test split — the model is tested on data it has never seen", bold_prefix="Validation:")
    _add_bullet(doc, " R² score (how well predictions match reality), MAE (average dollar error), MAPE (average percentage error)", bold_prefix="Metrics Reported:")
    _add_bullet(doc, " UNDERVALUED, FAIR VALUE, or OVERVALUED — with percentage discount/premium", bold_prefix="Output:")

    _add_para(doc, "")
    _add_para(doc, "Why this matters:", bold=True)
    _add_para(doc,
        "A traditional appraiser might take weeks and charge $5,000-$15,000 to produce a property "
        "valuation. While our ML model is NOT a replacement for a professional appraisal (and we're "
        "transparent about that), it provides an instant data-driven sanity check that catches obvious "
        "overpaying — or identifies potential bargains — before you commit resources to a full appraisal."
    )

    doc.add_page_break()

    # --- 4.2 ---
    _add_heading(doc, "4.2  Predictive Rent Growth (Polynomial Regression + Zillow ZORI Blending)", level=2)

    _add_para(doc, "What it does:", bold=True)
    _add_para(doc,
        "This model predicts how rents will grow over the next 7-10 years. Rent growth is the single "
        "most important driver of real estate investment returns, so getting this right is critical."
    )

    _add_para(doc, "The dual-data approach:", bold=True)
    _add_para(doc,
        "Most tools just let you type in a flat assumption like \"3% rent growth per year.\" That's "
        "unrealistic — rents don't grow at a constant rate. Our tool uses TWO real data sources and "
        "blends them together:"
    )

    _add_bullet(doc, " 60+ years of CPI Shelter data from the Federal Reserve (FRED). This captures the long-term national trend of how housing costs change over time.", bold_prefix="Data Source 1 — FRED CPI Shelter:")
    _add_para(doc, "")
    _add_bullet(doc, " City-level rent data from Zillow's Observed Rent Index (ZORI) covering 35 major metro areas. This captures LOCAL rent trends specific to the property's city.", bold_prefix="Data Source 2 — Zillow ZORI:")

    _add_para(doc, "")
    _add_para(doc, "The blending formula:", bold=True)
    _add_para(doc,
        "When Zillow data is available for the property's city, the model blends: 60% Zillow ZORI "
        "city-specific trend + 40% CPI national model. This gives you the best of both worlds — "
        "local accuracy with national stability. The ZORI forecast also includes mean-reversion "
        "(the assumption that extreme growth rates will gradually return toward the average), which "
        "prevents unrealistically optimistic or pessimistic projections."
    )

    _add_para(doc, "How the predictions feed into the financial model:", bold=True)
    _add_para(doc,
        "Instead of using a flat 3% growth rate for all 7 years, the AI generates VARIABLE year-by-year "
        "growth rates (e.g., 4.2% in Year 1, 3.8% in Year 2, 3.5% in Year 3...). These rates are "
        "injected directly into the pro forma financial model, making the 10-year projection far more "
        "realistic than any flat-rate assumption."
    )

    doc.add_page_break()

    # --- 4.3 ---
    _add_heading(doc, "4.3  Historical Backtest Validation", level=2)

    _add_para(doc, "What it does:", bold=True)
    _add_para(doc,
        "This is the \"prove it\" feature. Anyone can build a prediction model — but how do you know "
        "it actually works? The backtest takes historical data, hides the most recent portion, trains "
        "the model on the older data, and then checks how well it would have predicted the data that "
        "was hidden. It's like a time-travel test for the model."
    )

    _add_para(doc, "How it works:", bold=True)
    _add_bullet(doc, "Takes the full CPI Shelter historical dataset (60+ data points)")
    _add_bullet(doc, "Splits it 70/30 — trains on the first 70%, tests on the last 30%")
    _add_bullet(doc, "Compares what the model PREDICTED vs what ACTUALLY happened")
    _add_bullet(doc, "Reports three key accuracy metrics:")

    _add_para(doc, "")
    _add_table(doc,
        ["Metric", "What It Measures", "What \"Good\" Looks Like"],
        [
            ["MAE (Mean Absolute Error)", "Average size of prediction errors", "Below 0.5% = Strong"],
            ["RMSE (Root Mean Square Error)", "Penalizes large errors more heavily", "Below 0.5% = Strong"],
            ["Direction Accuracy", "Did we predict UP when it went UP?", "Above 70% = Strong"],
        ]
    )

    _add_para(doc, "")
    _add_para(doc, "Why this matters:", bold=True)
    _add_para(doc,
        "In the world of AI/ML, showing that your model works on historical data is the gold standard "
        "of credibility. This backtest gives you — and anyone reviewing your analysis — confidence "
        "that the rent predictions are based on a model with proven accuracy, not just wishful thinking. "
        "The results are displayed on the dashboard with a quality badge (STRONG, MODERATE, FAIR, or WEAK) "
        "and a chart showing actual vs. predicted values."
    )

    doc.add_page_break()

    # --- 4.4 ---
    _add_heading(doc, "4.4  Monte Carlo Simulation (1,000 Scenario Analysis)", level=2)

    _add_para(doc, "What it does:", bold=True)
    _add_para(doc,
        "This is perhaps the most powerful feature in the entire tool. Instead of producing a single "
        "\"best guess\" projection, the Monte Carlo simulation runs the ENTIRE financial model 1,000 "
        "times, each time with slightly different assumptions. The result is not one answer, but a "
        "probability distribution — essentially, it tells you: \"Here are all 1,000 possible futures "
        "for this investment, and here's how likely each outcome is.\""
    )

    _add_para(doc, "What it randomizes in each of the 1,000 runs:", bold=True)

    _add_table(doc,
        ["Assumption", "How It's Randomized", "Real-World Rationale"],
        [
            ["Rent Growth", "±30% of base rate", "Rents could grow faster or slower\nthan expected"],
            ["Occupancy", "±3 percentage points", "Vacancy rates fluctuate with\nmarket conditions"],
            ["Exit Cap Rate", "±50 basis points", "The sale price depends on market\nconditions at exit"],
            ["Expense Growth", "±25% of base rate", "Operating costs can surprise\nin either direction"],
        ]
    )

    _add_para(doc, "")
    _add_para(doc, "What you get:", bold=True)
    _add_bullet(doc, " \"75% chance your IRR exceeds 12%\" — this is infinitely more useful than a single number", bold_prefix="Probability statements:")
    _add_bullet(doc, " P5 (worst case), P25, P50 (median), P75, P95 (best case)", bold_prefix="Percentile breakdown:")
    _add_bullet(doc, " A visual chart showing the bell curve of all 1,000 outcomes, color-coded from red (losses) to green (strong returns)", bold_prefix="Distribution histogram:")
    _add_bullet(doc, " Probability of IRR > 8%, > 12%, > 15%, and probability of losing money (IRR < 0%)", bold_prefix="Risk metrics:")

    _add_para(doc, "")
    _add_para(doc, "Why this matters:", bold=True)
    _add_para(doc,
        "Professional investment firms like Blackstone, Brookfield, and Starwood ALL use Monte Carlo "
        "simulation in their underwriting. It's considered the gold standard for risk assessment because "
        "it accounts for uncertainty. A deal that shows 15% IRR in a single projection might look great — "
        "but if Monte Carlo reveals a 30% chance of losing money, you'd think twice. This tool brings "
        "that institutional-grade risk analysis to everyone."
    )

    doc.add_page_break()

    # --- 4.5 ---
    _add_heading(doc, "4.5  Natural Language Processing (NLP) for Lease Analysis", level=2)

    _add_para(doc, "What it does:", bold=True)
    _add_para(doc,
        "Upload a lease document (PDF), and the tool uses Google's Gemini AI to READ the entire lease "
        "and extract critical information automatically. This is Natural Language Processing — teaching "
        "a computer to understand human language in legal documents."
    )

    _add_para(doc, "What the AI extracts:", bold=True)
    _add_bullet(doc, "Tenant name, lease type (NNN, Gross, Modified Gross)")
    _add_bullet(doc, "Monthly rent, annual rent, lease term")
    _add_bullet(doc, "Escalation clauses (how rent increases are structured)")
    _add_bullet(doc, "Renewal options, security deposits, TI allowances")
    _add_bullet(doc, "CAM (Common Area Maintenance) charges")
    _add_bullet(doc, "Key clauses and risk flags")

    _add_para(doc, "")
    _add_para(doc, "The smart cross-check:", bold=True)
    _add_para(doc,
        "After extracting lease data, the tool automatically COMPARES the lease terms to the numbers "
        "you entered in the form. If the lease says rent is $1,500/month but you entered $1,200, the "
        "tool flags this discrepancy with a warning. This catches data entry errors and ensures your "
        "analysis matches the actual legal obligations."
    )

    _add_para(doc, "")
    _add_para(doc, "Multi-lease support:", bold=True)
    _add_para(doc,
        "For properties with multiple tenants (like an office building or strip mall), you can upload "
        "multiple lease PDFs. The AI analyzes each one individually AND produces a portfolio summary, "
        "totaling all rents and flagging risks across the entire tenant base."
    )

    doc.add_page_break()

    # --- 4.6 ---
    _add_heading(doc, "4.6  Multi-Signal AI Recommendation Engine", level=2)

    _add_para(doc, "What it does:", bold=True)
    _add_para(doc,
        "All of the AI/ML systems feed into a unified recommendation engine that produces a final "
        "verdict: STRONG BUY, BUY, HOLD/CONDITIONAL, or PASS. But unlike a simple thumbs-up/thumbs-down, "
        "it shows you exactly WHY it reached that conclusion by listing every signal."
    )

    _add_para(doc, "The six signals it weighs:", bold=True)

    _add_table(doc,
        ["Signal", "Source", "What It Evaluates"],
        [
            ["IRR Signal", "Financial Model", "Is the return above 8%? 12%? 15%?"],
            ["DSCR Signal", "Financial Model", "Can the property cover its debt\npayments?"],
            ["ML Valuation", "GradientBoosting Model", "Is the price fair, cheap, or\nexpensive?"],
            ["Rent Forecast", "Rent Predictor + ZORI", "Is the rent growth outlook\nstrong or weak?"],
            ["Monte Carlo", "1,000 Simulations", "What's the probability of\nachieving target returns?"],
            ["Lease Analysis", "NLP / Gemini AI", "Are there lease risks or\ndata discrepancies?"],
        ]
    )

    _add_para(doc, "")
    _add_para(doc,
        "Each signal contributes a positive or negative score. The final recommendation reflects the "
        "AGGREGATE view of all six systems. This multi-signal approach prevents any single metric from "
        "dominating the decision — just like a real investment committee that considers multiple perspectives."
    )

    doc.add_page_break()

    # --- 4.7 ---
    _add_heading(doc, "4.7  Interactive What-If Analysis", level=2)

    _add_para(doc, "What it does:", bold=True)
    _add_para(doc,
        "After the initial analysis is complete, the results page includes four interactive sliders "
        "that let you ask \"what if?\" questions and see the answers INSTANTLY — without re-running "
        "the entire analysis."
    )

    _add_para(doc, "The four sliders:", bold=True)
    _add_bullet(doc, " What if rents grow at 1% instead of 3%?", bold_prefix="Rent Growth (0-8%):")
    _add_bullet(doc, " What if I sell at a higher/lower cap rate?", bold_prefix="Exit Cap Spread (-50 to +100 bps):")
    _add_bullet(doc, " What if occupancy drops to 85%?", bold_prefix="Occupancy (75-99%):")
    _add_bullet(doc, " What if expenses grow faster at 5%?", bold_prefix="Expense Growth (0-6%):")

    _add_para(doc, "")
    _add_para(doc, "How it works technically:", bold=True)
    _add_para(doc,
        "When you drag a slider, the tool sends your new assumptions to the server via an API call "
        "(with 300ms debouncing to prevent overload). The server deep-copies the original deal, applies "
        "your modified assumptions, runs the FULL financial model with the new numbers, and returns "
        "updated metrics — all in under a second. The dashboard updates four key metrics (IRR, Equity "
        "Multiple, DSCR, Cash-on-Cash) and a NOI comparison chart in real-time."
    )

    _add_para(doc, "Why this matters:", bold=True)
    _add_para(doc,
        "In a real negotiation, you need to quickly evaluate: \"If I negotiate the price down 5%, "
        "what does that do to my returns?\" or \"If occupancy drops next year, can I still cover my "
        "debt?\" These sliders give you those answers instantly, making you a more informed and "
        "confident negotiator."
    )

    doc.add_page_break()

    # ==================== SECTION 5 ====================
    _add_heading(doc, "5. Real-World Data Integration", level=1)

    _add_para(doc,
        "The tool connects to five real-world data sources, ensuring every analysis is grounded in "
        "current market conditions rather than stale assumptions:",
        bold=True
    )

    _add_table(doc,
        ["Data Source", "What It Provides", "How It's Used"],
        [
            ["FRED (Federal Reserve\nEconomic Data)", "Treasury rates, CPI inflation,\nmortgage rates, housing starts,\nrental vacancy rates", "Cap rate estimation, rent prediction\ntraining, ML feature inputs, macro\nsignal calibration"],
            ["U.S. Census Bureau", "Median household income,\npopulation, median rent\nby metro area", "ML valuation features,\nmarket context for\nrecommendation"],
            ["Bureau of Labor\nStatistics (BLS)", "Local unemployment rate", "ML valuation feature,\neconomic health indicator"],
            ["Zillow ZORI", "City-level observed rent\nindex (35 metros)", "Rent growth prediction\nblending (60% weight)"],
            ["Google Gemini AI", "Natural language understanding\nof lease documents", "Lease term extraction,\nrisk flag identification"],
        ]
    )

    doc.add_page_break()

    # ==================== SECTION 6 ====================
    _add_heading(doc, "6. What the Tool Produces", level=1)

    _add_para(doc, "The tool generates a comprehensive package of deliverables:", bold=True)
    _add_para(doc, "")

    _add_heading(doc, "Interactive Web Dashboard", level=2)
    _add_bullet(doc, "Color-coded BUY/HOLD/PASS recommendation with signal breakdown")
    _add_bullet(doc, "Key return metrics: IRR, Equity Multiple, Cash-on-Cash, DSCR")
    _add_bullet(doc, "10-year pro forma table with before-tax and after-tax cash flows")
    _add_bullet(doc, "NOI growth chart and cash flow comparison chart")
    _add_bullet(doc, "Monte Carlo histogram and probability cards")
    _add_bullet(doc, "Backtest validation chart (actual vs. predicted)")
    _add_bullet(doc, "Interactive what-if sliders with live metric updates")
    _add_bullet(doc, "Sensitivity tables (exit cap, interest rate, rent growth, purchase price)")
    _add_bullet(doc, "ML valuation assessment with feature importance chart")
    _add_bullet(doc, "Lease analysis summary with risk flags")

    _add_para(doc, "")
    _add_heading(doc, "Downloadable Reports", level=2)

    _add_table(doc,
        ["Format", "Contents", "Best For"],
        [
            ["Excel Workbook\n(.xlsx)", "Multi-tab: Pro Forma, Sensitivity,\nMonte Carlo, Backtest, ML Valuation,\nLease Analysis", "Detailed number crunching,\ncustom analysis, auditing\nformulas"],
            ["Word Investment\nMemo (.docx)", "Institutional-style memo with\nall sections, tables, and\nrecommendation narrative", "Presenting to investors,\nlenders, or investment\ncommittees"],
            ["PDF Report\n(.pdf)", "Professional report with cover page,\nexecutive summary, all analysis\nsections, disclaimers", "Sharing, printing, archiving,\nemail attachments"],
        ]
    )

    _add_para(doc, "")
    _add_heading(doc, "Deal Comparison Mode", level=2)
    _add_para(doc,
        "Save multiple deals and view them side-by-side on a comparison page. The tool automatically "
        "color-codes the best values in each category (green = best), making it easy to see which "
        "deal offers the strongest returns, lowest risk, or best value. An NOI growth comparison "
        "chart shows how the properties' income trajectories compare over time."
    )

    doc.add_page_break()

    # ==================== SECTION 7 ====================
    _add_heading(doc, "7. Who Benefits From This Tool?", level=1)

    _add_table(doc,
        ["User", "How They Benefit"],
        [
            ["Real Estate Investors", "Get institutional-quality analysis before committing millions.\nIdentify undervalued properties and avoid overpaying."],
            ["Commercial Brokers", "Quickly underwrite deals for clients. Generate professional\nreports that build credibility and close deals faster."],
            ["Finance/RE Students", "Learn how professional underwriting works with a real tool.\nSee how AI/ML enhances traditional financial analysis."],
            ["Lenders & Banks", "Evaluate loan applications with data-driven risk metrics.\nMonte Carlo shows probability of borrower cash flow issues."],
            ["Small Investors", "Access the same analytical power as billion-dollar firms.\nNo expensive software licenses needed."],
            ["Professors & Researchers", "Demonstrate real-world application of GradientBoosting,\nPolynomial Regression, Monte Carlo, and NLP in finance."],
        ]
    )

    doc.add_page_break()

    # ==================== SECTION 8 ====================
    _add_heading(doc, "8. Technical Architecture Summary", level=1)

    _add_para(doc, "For the technically curious, here's how the system is built:", bold=True)
    _add_para(doc, "")

    _add_table(doc,
        ["Component", "Technology", "Details"],
        [
            ["Web Framework", "Flask (Python)", "Lightweight, production-ready with\ngunicorn deployment"],
            ["ML Valuation", "scikit-learn\nGradientBoostingRegressor", "18 features, 800 samples, 250\nestimators, 80/20 train/test split"],
            ["Rent Prediction", "scikit-learn\nPolynomialFeatures +\nLinearRegression", "Degree-2 polynomial on CPI Shelter\ndata, blended with Zillow ZORI"],
            ["Monte Carlo", "NumPy", "1,000 iterations, 4 randomized\nassumptions, full pro forma per run"],
            ["Backtest", "scikit-learn + NumPy", "70/30 time-based split, MAE/RMSE/\ndirection accuracy metrics"],
            ["Lease NLP", "Google Gemini API", "Structured JSON extraction from\nunstructured PDF documents"],
            ["Data APIs", "FRED, Census, BLS", "Real-time economic data via REST\nAPIs with caching"],
            ["Zillow Data", "Local CSV (ZORI)", "35 metro areas, semi-annual rent\nindex 2019-2025"],
            ["Excel Generation", "openpyxl", "Multi-tab workbook with charts"],
            ["Word Generation", "python-docx", "Institutional investment memo format"],
            ["PDF Generation", "fpdf2", "Professional report with cover page"],
            ["Frontend", "Bootstrap 5 + Chart.js", "Dark theme, responsive, interactive\ncharts and sliders"],
            ["Deployment", "Render (cloud)", "Auto-deploy from GitHub, free tier"],
        ]
    )

    doc.add_page_break()

    # ==================== SECTION 9 ====================
    _add_heading(doc, "9. Why This Matters — The Bigger Picture", level=1)

    _add_para(doc,
        "The real estate industry manages over $30 trillion in assets in the United States alone. "
        "Yet the tools used to evaluate these investments have barely changed in decades. Most "
        "professionals still rely on static Excel spreadsheets, manual research, and gut feelings.",
        size=12
    )

    _add_para(doc,
        "This tool represents a fundamental shift: the application of modern AI and machine learning "
        "to an industry that desperately needs it. Here's what makes it special:",
        size=12
    )

    _add_para(doc, "")

    _add_heading(doc, "Democratization of Institutional-Quality Analysis", level=2)
    _add_para(doc,
        "Features like Monte Carlo simulation, ML-based valuation, and NLP lease analysis were "
        "previously available only to firms with teams of data scientists and six-figure software "
        "budgets. This tool makes them available to anyone, for free."
    )

    _add_heading(doc, "Data-Driven Decision Making", level=2)
    _add_para(doc,
        "Every recommendation is backed by real data from federal agencies, validated ML models, "
        "and 1,000-scenario simulations. It removes emotion from what is often a highly emotional "
        "decision — buying property — and replaces it with evidence."
    )

    _add_heading(doc, "Transparency and Honesty", level=2)
    _add_para(doc,
        "The tool is transparent about its limitations. The ML model clearly states it trains on "
        "synthetic data and is NOT a substitute for professional appraisals. The backtest proves "
        "the rent predictor's accuracy. The Monte Carlo shows the full range of possible outcomes, "
        "including the probability of losing money. This kind of intellectual honesty builds trust."
    )

    _add_heading(doc, "Speed and Efficiency", level=2)
    _add_para(doc,
        "What used to take a team of analysts 2-3 weeks now takes 30 seconds. This doesn't just "
        "save time — it changes how investors work. You can evaluate 50 deals in a day instead of "
        "2, dramatically increasing the chances of finding the best opportunity."
    )

    _add_para(doc, "")
    _add_para(doc, "")

    # Closing
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = closing.add_run("— End of Document —")
    run.font.size = Pt(14)
    run.font.color.rgb = GOLD
    run.italic = True

    _add_para(doc, "")

    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(
        "This tool is designed for educational and professional use. ML predictions are based on "
        "synthetic training data and should supplement — not replace — professional due diligence. "
        "Always consult qualified professionals before making investment decisions."
    )
    run.font.size = Pt(9)
    run.font.color.rgb = LIGHT_GRAY
    run.italic = True

    # Save
    doc.save(OUTPUT_PATH)
    print(f"\nDocument saved to: {OUTPUT_PATH}")
    print(f"File size: {os.path.getsize(OUTPUT_PATH) / 1024:.1f} KB")
    return OUTPUT_PATH


if __name__ == "__main__":
    generate()
